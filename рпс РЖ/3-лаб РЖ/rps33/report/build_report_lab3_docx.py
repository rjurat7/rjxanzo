# -*- coding: utf-8 -*-
"""Generate Word report for Lab 3 (rps33) — aligned with methodical requirements and codebase."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt


def add_para(doc, text, first_indent_cm=1.25, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(14)
    r.bold = bold
    p.paragraph_format.first_line_indent = Cm(first_indent_cm)
    p.paragraph_format.space_after = Pt(6)
    return p


def h(doc, text, level=1):
    hd = doc.add_heading(text, level=level)
    for run in hd.runs:
        run.font.name = "Times New Roman"
    return hd


def main():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)
    try:
        doc.styles["Normal"].font.name = "Times New Roman"
        doc.styles["Normal"].font.size = Pt(14)
    except Exception:
        pass

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(
        "МИНОБРНАУКИ РОССИИ\n"
        "федеральное государственное бюджетное образовательное учреждение высшего образования\n"
        "«Санкт-Петербургский государственный технологический институт (технический университет)»"
    )
    r.font.size = Pt(12)

    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = t2.add_run(
        "Отчёт по лабораторной работе № 3\n"
        "Вариант № 22\n"
        "Тема: «Изучение реализации графических пользовательских интерфейсов и интеграции базы данных "
        "для веб-приложений»"
    )
    r2.font.size = Pt(14)
    r2.bold = True
    doc.add_paragraph()

    h(doc, "СОДЕРЖАНИЕ", 1)
    for line in [
        "1. Цель работы",
        "2. Задание",
        "3. Описание используемых технологий (стек БД и GUI)",
        "4. Структура базы данных",
        "5. Результаты выполнения интеграционных тестов (п. 3 задания)",
        "6. Описание графического интерфейса и эвристики UI/UX",
        "7. Выводы",
    ]:
        add_para(doc, line, first_indent_cm=0)

    doc.add_page_break()

    h(doc, "1. Цель работы", 1)
    add_para(
        doc,
        "Разработка графического пользовательского интерфейса к программе, интеграция базы данных "
        "для хранения исходных и отсортированных целочисленных массивов, реализация интеграционных тестов "
        "с отдельной тестовой базой, отражённых в отдельном консольном проекте.",
    )

    h(doc, "2. Задание", 1)
    add_para(
        doc,
        "В репозитории создать ветку develop-lab3; развернуть СУБД; спроектировать схему хранения массивов; "
        "реализовать тесты вставки (100/1000/10 000), выгрузки и сортировки 100 массивов (3 прогона на объёме "
        "100/1000/10 000 записей), очистки БД (3 прогона). Тесты — отдельно от GUI, с отдельной БД. "
        "Настольное приложение: регистрация/вход, ввод с клавиатуры и генерация массива, сортировка "
        "алгоритмом лаб. №2, сохранение, просмотр истории. Отчёт — по структуре из методички.",
    )

    h(doc, "3. Описание используемых технологий", 1)
    add_para(doc, "3.1. База данных — SQLite (встраиваемая СУБД)", bold=True)
    add_para(
        doc,
        "В проекте rps33 используется библиотека SQLite (файл БД sorting.db). Достоинства: не требуется "
        "отдельный сервер, одна портативная БД-файл, SQL, индексы и внешние ключи, кросс-платформенность, "
        "низкая стоимость внедрения. Недостатки: ориентирован на локальные сценарии, не решает задачи "
        "кластерного высокодоступного сервера «из коробки».",
    )
    add_para(doc, "3.2. Графический интерфейс — C++/CLI, Windows Forms", bold=True)
    add_para(
        doc,
        "Интерфейс — класс MainForm, WinForms-контролы, вызовы нативного кода (Database, ShakerSort) из управляемого. "
        "Плюсы: быстрая разработка, типичный стек для учебного desktop. Минусы: внешний вид и темизация слабее, "
        "чем у WPF/веб; привязка к экосистеме Windows.",
    )

    h(doc, "4. Структура базы данных", 1)
    add_para(
        doc,
        "Таблица users: id (INTEGER, PK, AUTOINCREMENT), login (TEXT, UNIQUE), password (TEXT), created_at. "
        "Таблица arrays: id (PK), user_id (FK → users, ON DELETE CASCADE), original (TEXT — массив как строка), "
        "sorted (TEXT), created_at. Массивы в полях original/sorted кодируются перечислением чисел через пробел; "
        "сериализация/десериализация в классе Database (vectorToString, stringToVector).",
    )

    h(doc, "5. Результаты выполнения интеграционных тестов", 1)
    add_para(
        doc,
        "Отдельный проект rps33_tests, исполняемый файл запускается независимо от rps33.exe. Используется "
        "отдельный файл БД sorting_test.db. В main.cpp: тест 1 — пакетная вставка test_user/случайных массивов; "
        "тест 2 — трижды getArraysForSorting(100) + шейкерная сортировка, вывод total_s и avg_s; тест 3 — "
        "очистка (clearTestData) трижды с повторным наполнением между прогонами. Запуск: собрать rps33_tests, "
        "выполнить, вставить в отчёт фактические строки вывода (time_s, total_s, avg_s).",
    )

    h(doc, "6. Описание графического интерфейса, достоинства, недостатки, UI/UX (п. 4 задания)", 1)
    h(doc, "6.1. Достоинства (реализация MainForm, WinForms)", 2)
    for line in [
        "Логичная группировка: три панели groupAuth, groupArray, groupHistory (сверху вниз: вход → массив → история).",
        "Простота и понятность: подписи полей, однозначные названия кнопок.",
        "Предзаполненное поле txtArray (пример массива для быстрого теста).",
        "Понятные надписи: «Войти», «Регистрация», «Сгенерировать», «Сортировать», «Сохранить в БД».",
        "Статусная строка lblStatus (состояние БД, результаты операций, цвета).",
        "Всплывающие MessageBox (ошибка ввода, вход не выполнен, успех сохранения).",
        "Естественный порядок: авторизация → ввод/генерация/сортировка → сохранение → история.",
        "Сохранение без авторизации заблокировано логически: при пустом currentUsername показ предупреждения, запись в БД не выполняется.",
    ]:
        add_para(doc, line, first_indent_cm=1.0)

    h(doc, "6.2. Недостатки (учебная версия)", 2)
    for line in [
        "Нет кастомных тем/стилей на уровне дизайн-системы.",
        "Нет горячих клавиш (Ctrl+S, F1 привязан к справке не везде).",
        "Нет контекстного меню в listHistory.",
        "Нет экспорта истории в файл.",
        "Нет поиска/фильтра по истории.",
        "Пароли в SQLite не хэшируются.",
        "Нет блокировки учётной записи после неудачных попыток входа.",
    ]:
        add_para(doc, line, first_indent_cm=1.0)

    h(doc, "6.3. Соответствие эвристикам (кратко)", 2)
    add_para(
        doc,
        "Видимость статуса: lblStatus. Соответствие миру и согласованность: привычные подписи и сценарий. "
        "Минимализм: три группы без лишних экранов. Ошибки: MessageBox, подсказки ToolTip на кнопках. "
        "Справка: кнопка «Справка» (btnHelp).",
    )

    h(doc, "6.4. Структура окна (имена контролов)", 2)
    add_para(
        doc,
        "1) Группа «Авторизация» (groupAuth, lblAuthTitle): txtLogin, txtPassword, chkShowPassword, btnLogin, "
        "btnRegister, btnLogout, btnClearHistory. "
        "2) Группа «Работа с массивом» (groupArray, lblArrayTitle): txtArray (RichTextBox), btnGenerate, "
        "btnSort, btnSave. 3) Группа «История» (groupHistory, lblHistoryTitle): listHistory, подписи. "
        "4) Внизу: lblStatus, btnHelp.",
    )

    h(doc, "7. Выводы", 1)
    add_para(
        doc,
        "Создано десктопное приложение с WinForms и SQLite, реализованы требуемые сценарии работы с массивами и "
        "пользователями, отдельный набор интеграционных тестов с отдельной тестовой БД. Для учебного проекта "
        "указанные технологии достаточны; для «боевой» системы потребуется, например, хэширование паролей и усиление "
        "безопасности.",
    )

    doc.add_page_break()
    h(doc, "Сведения", 1)
    tbl = doc.add_table(rows=5, cols=2)
    rows = [
        ("Исполнитель", "Рузматов Ж.Т."),
        ("Проект", "rps33 (Visual Studio)"),
        ("БД (приложение)", "sorting.db"),
        ("БД (тесты)", "sorting_test.db"),
        ("Проверяющий", "Дамрин А.О."),
    ]
    for i, (a, b) in enumerate(rows):
        tbl.rows[i].cells[0].text = a
        tbl.rows[i].cells[1].text = b

    out = Path(__file__).resolve().parent / "Отчет_ЛР3_РПС_Рузматов_Ж_обновлено.docx"
    doc.save(str(out))
    print("Saved:", out)

    # Copy to Desktop
    try:
        import shutil
        desk = Path.home() / "Desktop" / "Отчет_ЛР3_РПС_Рузматов_Ж_обновлено.docx"
        shutil.copy2(out, desk)
        print("Copied to:", desk)
    except Exception as e:
        print("Desktop copy skip:", e)


if __name__ == "__main__":
    main()
