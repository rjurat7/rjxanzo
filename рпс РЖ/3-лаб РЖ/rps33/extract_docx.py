from pathlib import Path
from docx import Document
def dump_docx(p, out):
    doc = Document(str(p))
    lines = [x.strip() for x in (par.text for par in doc.paragraphs) if x.strip()]
    for t in doc.tables:
        for row in t.rows:
            c = " | ".join((c.text or "").strip() for c in row.cells if (c.text or "").strip())
            if c: lines.append(c)
    Path(out).write_text("\n".join(lines), encoding="utf-8")
    print(len(lines), out)
base = Path(r"d:\rps-main\rps33")
base.mkdir(parents=True, exist_ok=True)
req = Path(r"c:\Users\DILSHOD\Downloads\Лабораторная_работа_3 (1).docx")
rep = Path(r"c:\Users\DILSHOD\Desktop\Отчет РПС 3 Рузматов Ж.docx")
if req.exists():
    dump_docx(req, base / "_req_lab3.txt")
if rep.exists():
    dump_docx(rep, base / "_user_report_extracted.txt")
